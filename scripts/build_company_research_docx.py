"""Build a professional company-research DOCX from a workflow JSON artifact."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

WIDTH = 9360
INDENT = 120
NAVY, BLUE, DARK_BLUE = "0B2545", "2E74B5", "1F4D78"
PALE, GRAY, MID = "E8EEF5", "F2F4F7", "6B7280"
INK, WHITE, RED, AMBER, GREEN = "111827", "FFFFFF", "9B1C1C", "A16207", "1F6B45"
FONT, CN_FONT = "Calibri", "Microsoft YaHei"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def font(run, size=11, *, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), CN_FONT)
    run.font.size = Pt(size)
    run.bold, run.italic = bold, italic
    run.font.color.rgb = rgb(color)
    return run


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=90, bottom=90, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        tc_pr.append(node)
    for name, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        child = node.find(qn(f"w:{name}"))
        if child is None:
            child = OxmlElement(f"w:{name}")
            node.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def geometry(table, widths: list[int]):
    if sum(widths) != WIDTH:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr
    for tag, value in (("tblW", WIDTH), ("tblInd", INDENT)):
        node = pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(value))
        grid.append(node)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def cell_text(cell, value: str, *, bold=False, color=INK, size=9, align=None):
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    font(p.add_run(value), size, bold=bold, color=color)


def table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    for cell, label in zip(result.rows[0].cells, headers, strict=True):
        shade(cell, NAVY)
        cell_text(
            cell,
            label,
            bold=True,
            color=WHITE,
            size=8.7,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for values in rows:
        cells = result.add_row().cells
        for index, value in enumerate(values):
            cell_text(
                cells[index],
                value,
                size=8.7,
                align=WD_ALIGN_PARAGRAPH.CENTER
                if index > 0
                else WD_ALIGN_PARAGRAPH.LEFT,
            )
    geometry(result, widths)
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    result.rows[0]._tr.get_or_add_trPr().append(header)
    for row in result.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    return result


def numbering(doc: Document, bullet=True) -> int:
    root = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    aid, nid = max(abstract_ids, default=0) + 1, max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(aid))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("start", "1"),
        ("numFmt", "bullet" if bullet else "decimal"),
        ("lvlText", "•" if bullet else "%1."),
        ("lvlJc", "left"),
    ):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.append(p_pr)
    abstract.append(level)
    first_num = root.find(qn("w:num"))
    root.insert(
        list(root).index(first_num), abstract
    ) if first_num is not None else root.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(nid))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(aid))
    num.append(ref)
    root.append(num)
    return nid


def list_item(doc, text: str, num_id: int):
    p = doc.add_paragraph()
    pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([level, num])
    pr.append(num_pr)
    font(p.add_run(text))


def body(doc, text: str, lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if lead:
        font(p.add_run(lead), bold=True, color=NAVY)
    font(p.add_run(text))


def citation(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(text), 8.3, color=MID, italic=True)


def callout(doc, label: str, text: str, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    edge = OxmlElement("w:left")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), "18")
    edge.set(qn("w:space"), "7")
    edge.set(qn("w:color"), color)
    borders.append(edge)
    pr.append(borders)
    font(p.add_run(label + "  "), bold=True, color=color)
    font(p.add_run(text))


def page_break_heading(doc, value: str):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(value, 1)


def setup(doc: Document):
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = (
        section.right_margin
    ) = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 14, MID, 0, 18),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(
        hp.add_run("QUANT RESEARCH AGENT  |  COMPANY RESEARCH"),
        8.5,
        bold=True,
        color=MID,
    )
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(fp.add_run("公开信息研究底稿  |  "), 8, color=MID)
    run = fp.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def chart(artifact: dict, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1500, 650), "white")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 28)
    bold = ImageFont.truetype("C:/Windows/Fonts/msyhbd.ttc", 38)
    draw.text(
        (80, 35), "比亚迪相对估值：目标公司 vs 同业中位数", fill="#0B2545", font=bold
    )
    metrics = next(
        x
        for x in artifact["analysis_bundle"]["artifacts"]
        if x["method"] == "relative_valuation"
    )["metrics"]
    labels = [("PE TTM", "pe_ttm"), ("PB MRQ", "pb_mrq"), ("PS TTM", "ps_ttm")]
    colors = ["#2E74B5", "#9CA3AF"]
    for idx, (label, key) in enumerate(labels):
        y = 170 + idx * 145
        target = metrics[key]["target"]
        peer = metrics[key]["peer_median"]
        draw.text((80, y + 15), label, fill="#374151", font=regular)
        scale = 850 / max(target, peer)
        for j, (name, value) in enumerate((("比亚迪", target), ("同业中位数", peer))):
            yy = y + j * 50
            draw.rounded_rectangle(
                (300, yy, 300 + int(value * scale), yy + 34), radius=8, fill=colors[j]
            )
            draw.text((1180, yy), f"{name} {value:.2f}x", fill="#111827", font=regular)
    image.save(path)


def build(input_path: Path, output_path: Path, chart_path: Path) -> Path:
    artifact = json.loads(input_path.read_text(encoding="utf-8"))
    report = artifact["company_research_report"]
    filing = artifact["analysis_context"]["company_filing_analysis"]
    extraction = artifact["analysis_context"]["company_filing_extraction"]
    valuation = next(
        x
        for x in artifact["analysis_bundle"]["artifacts"]
        if x["method"] == "relative_valuation"
    )["metrics"]
    company = artifact["analysis_context"]["company_data"]
    market = {x["name"]: x["value"] for x in company["market_metrics"]}
    financial = {x["name"]: x["value"] for x in company["financial_metrics"]}
    chart(artifact, chart_path)
    doc = Document()
    setup(doc)
    bullets = numbering(doc, True)
    numbers = numbering(doc, False)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("上市公司深度研究"), 11, bold=True, color=AMBER)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("比亚迪：全球化增长与现金转化的再平衡"), 30, bold=True, color=NAVY)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("业务结构、盈利质量、资本强度与相对估值研究"), 14, color=MID)
    table(
        doc,
        ["项目", "内容"],
        [
            ["证券代码", report["security_code"]],
            ["研究截止日", report["as_of_date"]],
            ["信息边界", "巨潮资讯年报页级证据 + BaoStock 公开指标"],
            ["评审结论", "APPROVED / 公开信息研究可交付"],
        ],
        [2100, 7260],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        p.add_run("本报告为研究作品，不构成投资建议、估值承诺或交易指令。"),
        9.5,
        italic=True,
        color=MID,
    )
    doc.add_page_break()

    doc.add_heading("核心观点", 1)
    callout(
        doc,
        "BOTTOM LINE",
        "比亚迪的增长引擎正由国内规模扩张转向汽车主业深化与海外放量，但2025年经营现金流下降、投资净流出扩大，同时估值显著高于所选同业中位数。下一阶段的关键不是证明增长是否存在，而是验证海外增长、毛利率与现金转化能否同时兑现。",
    )
    table(
        doc,
        ["关键指标", "当前读数", "研究含义"],
        [
            ["2025年营收增速", "+3.46%", "收入保持增长，但需与利润及现金流联合判断"],
            ["汽车业务收入占比", "80.68%", "核心业务集中度进一步提高"],
            ["境外收入增速 / 占比", "+40.05% / 38.65%", "海外成为主要增量来源"],
            ["经营现金流净额变动", "-55.69%", "现金转化是当前首要验证点"],
            [
                "PE / PB / PS溢价",
                f"{valuation['pe_ttm']['premium_discount']:.0%} / {valuation['pb_mrq']['premium_discount']:.0%} / {valuation['ps_ttm']['premium_discount']:.0%}",
                "市场已计入较高增长与质量预期",
            ],
        ],
        [2500, 2200, 4660],
    )
    citation(doc, "来源：2025年年报第29、30、40页；BaoStock截至2026-08-07的市场指标。")
    doc.add_heading("三项判断", 2)
    for text in [
        "业务结构更聚焦汽车，海外收入成为第二增长曲线；但国内收入下降与价格竞争意味着规模增长并不自动转化为利润改善。",
        "经营现金流591亿元，同比下降55.69%；投资活动净流出约1,975亿元，外部融资对扩张的支撑显著上升。",
        "相对上汽集团、长城汽车的简单中位数，比亚迪PE、PB、PS均有较高溢价，后续回报对盈利兑现与现金流改善更敏感。",
    ]:
        list_item(doc, text, bullets)
    doc.add_heading("研究结论与优先监测", 2)
    table(
        doc,
        ["优先级", "议题", "需要看到的证据"],
        [
            ["P1", "现金转化与资本强度", "经营现金流降幅收窄；资本开支回报路径清晰"],
            ["P1", "汽车业务毛利率", "价格竞争下毛利率企稳，产品结构升级有效"],
            ["P2", "海外本地化", "境外增长延续，贸易壁垒对交付与成本的扰动可控"],
            ["P2", "估值消化", "盈利增速与现金流支持当前相对溢价"],
        ],
        [1200, 3000, 5160],
    )

    page_break_heading(doc, "1. 业务模式与增长结构")
    body(doc, filing["business_model"])
    callout(
        doc,
        "研究解读",
        "多元业务提供技术与供应链协同，但汽车业务已贡献超过八成收入，集团基本面越来越取决于汽车的销量、单车盈利、海外扩张与资本开支效率。",
        BLUE,
    )
    doc.add_heading("1.1 分部表现", 2)
    table(
        doc,
        ["分部", "2025年收入", "同比", "毛利率", "判断"],
        [
            [
                "汽车及相关产品",
                "6,486亿元",
                "+5.06%",
                "20.49%",
                "增长核心，毛利率同比下降1.82个百分点",
            ],
            ["手机部件及组装等", "1,552亿元", "-2.74%", "6.29%", "收入与毛利率均承压"],
        ],
        [2200, 1650, 1200, 1300, 3010],
    )
    citation(
        doc,
        "页级证据：CNINFO-2AFABDBD2F-P029-segment_information、P030-segment_information。",
    )
    doc.add_heading("1.2 地域结构", 2)
    body(
        doc,
        "境外收入约3,107亿元，同比增长40.05%，占比由28.55%升至38.65%；中国地区收入同比下降11.17%。海外已从补充市场转为影响集团增速与估值叙事的核心变量。",
    )
    citation(doc, "页级证据：2025年年报第29-30页。")

    page_break_heading(doc, "2. 财务质量与现金流")
    callout(
        doc,
        "关键矛盾",
        "收入增长与经营现金流显著下降并存。单凭利润或营收无法判断扩张质量，必须继续拆解库存、应收、供应商结算与资本开支。",
        RED,
    )
    table(
        doc,
        ["现金流项目", "2025年", "同比变化", "主要解释"],
        [
            ["经营活动净额", "591亿元", "-55.69%", "购买商品、接受劳务支付现金增加"],
            ["投资活动净额", "-1,975亿元", "净流出扩大52.97%", "购建长期资产支付增加"],
            ["筹资活动净额", "+1,046亿元", "+1,118.88%", "H股配售及借款、发债增加"],
        ],
        [2300, 1800, 1900, 3360],
    )
    citation(doc, "页级证据：CNINFO-2AFABDBD2F-P040-cash_flow（2025年年报第40页）。")
    doc.add_heading("2.1 最新公开财务指标（2026Q1）", 2)
    table(
        doc,
        ["指标", "读数", "解释边界"],
        [
            [
                "ROE（BaoStock口径）",
                f"{financial['profit.roeAvg']:.2%}",
                "季度口径，不与全年直接类比",
            ],
            ["净利率", f"{financial['profit.npMargin']:.2%}", "反映阶段性盈利水平"],
            [
                "资产负债率",
                f"{financial['balance.liabilityToAsset']:.2%}",
                "需结合有息负债与经营负债拆解",
            ],
            [
                "流动比率 / 速动比率",
                f"{financial['balance.currentRatio']:.2f}x / {financial['balance.quickRatio']:.2f}x",
                "仅作流动性筛查",
            ],
            [
                "CFO / 净利润",
                f"{financial['cash_flow.CFOToNP']:.2f}x",
                "需结合季度季节性解释",
            ],
        ],
        [2900, 2200, 4260],
    )
    citation(
        doc,
        "来源：BaoStock，指标发布日期2026-04-29；不替代对一季报原始报表的逐项复核。",
    )

    page_break_heading(doc, "3. 竞争位置、管理层重点与风险")
    body(doc, filing["competitive_position"])
    doc.add_heading("3.1 管理层优先事项", 2)
    for item in filing["management_priorities"]:
        list_item(doc, item, bullets)
    citation(doc, "页级证据：2025年年报第48页及相关管理层讨论摘录。")
    doc.add_heading("3.2 二维风险优先级", 2)
    table(
        doc,
        ["影响 / 紧迫度", "高紧迫度", "中紧迫度"],
        [
            ["高影响", "现金转化与资本开支；汽车毛利率", "海外贸易壁垒；估值溢价兑现"],
            ["中影响", "国内价格竞争", "手机部件业务承压；原材料波动"],
        ],
        [2100, 3630, 3630],
    )
    doc.add_heading("3.3 风险链条", 2)
    for item in filing["risks"]:
        list_item(doc, item, bullets)
    callout(
        doc,
        "可信边界",
        "年报确认了风险存在与管理层应对方向；关于库存、应收或流动性恶化的表述属于研究假设，尚需原始报表科目和后续季度数据验证。",
        AMBER,
    )

    page_break_heading(doc, "4. 相对估值与市场定价")
    body(
        doc,
        "相对估值只回答市场给了多少倍数，不直接回答公司值多少钱。本报告使用上汽集团与长城汽车的同步市场指标中位数作为窄口径参照，结果对同业选择高度敏感。",
    )
    doc.add_picture(str(chart_path), width=Inches(6.35))
    picture = doc.inline_shapes[-1]
    picture._inline.docPr.set("title", "比亚迪相对估值比较")
    picture._inline.docPr.set(
        "descr", "比亚迪与上汽集团、长城汽车同业中位数的PE、PB、PS柱状比较图"
    )
    citation(
        doc,
        "图1：比亚迪与所选同业中位数的PE、PB、PS比较；来源：BaoStock，截至2026-08-07。",
    )
    table(
        doc,
        ["指标", "比亚迪", "同业中位数", "溢价"],
        [
            [
                "PE TTM",
                f"{valuation['pe_ttm']['target']:.2f}x",
                f"{valuation['pe_ttm']['peer_median']:.2f}x",
                f"{valuation['pe_ttm']['premium_discount']:.1%}",
            ],
            [
                "PB MRQ",
                f"{valuation['pb_mrq']['target']:.2f}x",
                f"{valuation['pb_mrq']['peer_median']:.2f}x",
                f"{valuation['pb_mrq']['premium_discount']:.1%}",
            ],
            [
                "PS TTM",
                f"{valuation['ps_ttm']['target']:.2f}x",
                f"{valuation['ps_ttm']['peer_median']:.2f}x",
                f"{valuation['ps_ttm']['premium_discount']:.1%}",
            ],
        ],
        [2500, 2200, 2400, 2260],
    )
    callout(
        doc,
        "估值判断",
        "较高溢价可由增长、技术与全球化预期解释，但也抬高了对利润率、现金流和海外执行的要求。若现金转化持续弱于收入增长，估值溢价的容错空间会下降。",
        BLUE,
    )
    body(
        doc,
        f"市场语境：截至2026-08-07，收盘价{market['close']:.2f}元，近一年回报{market['one_year_return']:.1%}，年化波动率{market['annualized_volatility']:.1%}，区间最大回撤{market['max_drawdown']:.1%}。这些数据用于风险语境，不构成目标价。",
    )

    page_break_heading(doc, "5. 行动路线与跟踪框架")
    table(
        doc,
        ["行动", "Owner", "Timeline", "KPI / 验证标准"],
        [
            [
                "建立现金转化桥接表",
                "财务研究",
                "0-1个月",
                "经营现金流差异可归因至库存、应收、应付和预收",
            ],
            [
                "拆解汽车毛利率驱动",
                "行业研究",
                "0-2个月",
                "形成价格、产品结构、成本的季度桥接",
            ],
            [
                "建立海外市场看板",
                "海外研究",
                "持续 / 月度",
                "销量、产能、关税、渠道与本地化进度同步更新",
            ],
            [
                "扩展可比公司池",
                "估值研究",
                "0-1个月",
                "按业务、增长、盈利和地域进行分层比较",
            ],
            [
                "设置研究更新触发器",
                "Research Manager",
                "持续",
                "季报、产销快报、重大贸易政策发布后自动重跑",
            ],
        ],
        [2600, 1500, 1600, 3660],
    )
    doc.add_heading("5.1 研究委员会待回答问题", 2)
    for text in [
        "经营现金流下降主要来自增长性营运资金投入，还是议价能力或回款质量变化？",
        "海外收入增长能否在关税、本地化投入和渠道扩张后维持可接受利润率？",
        "汽车业务毛利率下降是周期性价格竞争，还是结构性定价权弱化？",
        "当前估值溢价需要怎样的盈利增速和现金流改善才能被合理消化？",
    ]:
        list_item(doc, text, numbers)
    callout(
        doc,
        "委员会结论",
        artifact["company_research_review"]["overall_assessment"]
        + " 结论为条件性通过：后续季度数据变化将触发重新评审。",
    )

    page_break_heading(doc, "附录：证据索引与研究边界")
    rows = []
    for item in extraction["page_evidence"]:
        rows.append(
            [
                item["evidence_id"],
                f"第{item['page_number']}页",
                item["title"].split(" - ")[-1],
                item["source_name"],
            ]
        )
    table(doc, ["证据ID", "页码", "主题", "来源"], rows, [4600, 1100, 2400, 1260])
    citation(
        doc,
        f"原始年报：{extraction['title']}；PDF SHA-256：{extraction['sha256']}；共{extraction['page_count']}页。",
    )
    doc.add_heading("研究限制", 2)
    for item in report["limitations"]:
        list_item(doc, item, bullets)
    for text in [
        "DeepSeek仅在已抽取的年报页证据范围内生成解释；页码、指标、路由和评审由代码校验。",
        "同业样本仅包含上汽集团与长城汽车，不能代表完整汽车产业可比公司集合。",
        "本报告未建立DCF或目标价模型，亦未接入非公开经营信息、渠道访谈和供应链尽调。",
    ]:
        list_item(doc, text, bullets)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path.resolve()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--chart", type=Path, required=True)
    args = parser.parse_args()
    print(build(args.input, args.output, args.chart))


if __name__ == "__main__":
    main()
