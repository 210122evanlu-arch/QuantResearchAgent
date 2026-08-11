"""Build a polished DOCX from the evidence-validated DeepSeek narrative artifact."""

import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
DARK_GRAY = "374151"
GOLD = "B38A3E"
RED = "9B1C1C"
AMBER = "A16207"
GREEN = "1F6B45"
WHITE = "FFFFFF"
BLACK = "111827"
ASCII_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = ASCII_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ASCII_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = _rgb(color)
    if italic is not None:
        run.italic = italic


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], *, indent=TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_cell_text(cell, text: str, *, bold=False, color=BLACK, size=9.5, align=None):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text) if not paragraph.text else paragraph.runs[0]
    if paragraph.text and paragraph.runs:
        paragraph.runs[0].text = text
        run = paragraph.runs[0]
    _set_run_font(run, size=size, bold=bold, color=color)


def _set_paragraph_border(paragraph, color: str, size=12, space=6, side="left") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def _add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = ASCII_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = _rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    tokens = {
        "Title": (28, NAVY, 0, 8),
        "Subtitle": (14, DARK_GRAY, 0, 18),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = ASCII_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ASCII_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = _rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True


def _add_numbering(doc: Document, *, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if bullet else "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    level.extend([start, num_fmt, level_text, level_jc, p_pr])
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _add_list_item(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    run = paragraph.add_run(text)
    _set_run_font(run, size=11, color=BLACK)


def _add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.widow_control = True
    if bold_lead:
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, size=11, bold=True, color=NAVY)
    run = paragraph.add_run(text)
    _set_run_font(run, size=11, color=BLACK)


def _add_citation(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    _set_run_font(run, size=8.5, italic=True, color=MID_GRAY)


def _add_callout(doc: Document, label: str, text: str, *, color=NAVY) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.15
    _set_paragraph_border(paragraph, color, size=16, space=7)
    lead = paragraph.add_run(label + "  ")
    _set_run_font(lead, size=11, bold=True, color=color)
    body = paragraph.add_run(text)
    _set_run_font(body, size=11, color=BLACK)


def _add_section_heading(doc: Document, text: str, *, page_break=True) -> None:
    if page_break:
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading(text, level=1)


def _configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(
        "FINANCIAL RESEARCH & ADVISORY  |  PUBLIC-INFORMATION RISK REVIEW"
    )
    _set_run_font(run, size=8.5, bold=True, color=MID_GRAY)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Confidential working paper  |  ")
    _set_run_font(run, size=8, color=MID_GRAY)
    _add_field(paragraph, "PAGE")


def _add_cover(doc: Document, report: dict, artifact: dict) -> None:
    for _ in range(5):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("上市公司风险与战略咨询")
    _set_run_font(run, size=11, bold=True, color=GOLD)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report["title"])
    _set_run_font(run, size=28, bold=True, color=NAVY)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(report["subtitle"])
    _set_run_font(run, size=14, color=DARK_GRAY)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    _set_table_geometry(meta, [2520, 6840])
    rows = [
        ("分析对象", "比亚迪股份有限公司（002594.SZ）"),
        ("分析截止日", artifact["as_of_date"]),
        ("报告类型", "公开信息风险咨询 | 管理层讨论稿"),
        ("生成方式", f"证据约束分析 + Debate Gate + {artifact['model']}咨询写作"),
    ]
    for index, (label, value) in enumerate(rows):
        _set_cell_shading(meta.rows[index].cells[0], LIGHT_GRAY)
        _set_cell_text(
            meta.rows[index].cells[0], label, bold=True, color=NAVY, size=9.5
        )
        _set_cell_text(meta.rows[index].cells[1], value, color=BLACK, size=9.5)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("仅供研究与管理讨论，不构成投资建议、审计意见或法律意见")
    _set_run_font(run, size=9.5, italic=True, color=MID_GRAY)
    doc.add_page_break()


def _build_charts(artifact: dict, output_dir: Path) -> tuple[Path, Path, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    finance_path = output_dir / "byd_financial_change.png"
    market_path = output_dir / "byd_market_rebased.png"
    risk_matrix_path = output_dir / "byd_risk_priority_matrix.png"

    width, height = 1500, 620
    font_path = Path("C:/Windows/Fonts/calibri.ttf")
    bold_path = Path("C:/Windows/Fonts/calibrib.ttf")
    font = ImageFont.truetype(str(font_path), 28)
    small = ImageFont.truetype(str(font_path), 24)
    title_font = ImageFont.truetype(str(bold_path), 38)

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (90, 35), "BYD 2025 selected financial changes", fill="#0B2545", font=title_font
    )
    labels = ["Revenue", "Net profit", "Operating CF", "Equity"]
    values = [3.46, -18.97, -55.69, 32.94]
    baseline = 310
    draw.line((100, baseline, 1420, baseline), fill="#6B7280", width=2)
    for index, (label, value) in enumerate(zip(labels, values, strict=True)):
        center = 250 + index * 330
        bar_width = 150
        magnitude = int(abs(value) * 4.1)
        top = baseline - magnitude if value >= 0 else baseline
        bottom = baseline if value >= 0 else baseline + magnitude
        color = f"#{BLUE}" if value >= 0 else f"#{RED}"
        draw.rounded_rectangle(
            (center - bar_width // 2, top, center + bar_width // 2, bottom),
            radius=8,
            fill=color,
        )
        value_y = top - 38 if value >= 0 else bottom + 8
        value_text = f"{value:+.2f}%"
        value_box = draw.textbbox((0, 0), value_text, font=font)
        draw.text(
            (center - (value_box[2] - value_box[0]) / 2, value_y),
            value_text,
            fill="#111827",
            font=font,
        )
        label_box = draw.textbbox((0, 0), label, font=small)
        draw.text(
            (center - (label_box[2] - label_box[0]) / 2, 555),
            label,
            fill="#374151",
            font=small,
        )
    image.save(finance_path, quality=95)

    raw_series = artifact["source_context"]["market"]["rebased_series"]
    dates = [datetime.fromisoformat(item["date"]) for item in raw_series]
    byd = [float(item["byd_rebased"]) for item in raw_series]
    csi300 = [float(item["csi300_rebased"]) for item in raw_series]
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((90, 35), "BYD versus CSI 300", fill="#0B2545", font=title_font)
    left, top, right, bottom = 110, 125, 1425, 520
    all_values = byd + csi300
    low, high = min(all_values), max(all_values)
    padding = max((high - low) * 0.08, 2)
    low -= padding
    high += padding

    def point(index: int, value: float) -> tuple[int, int]:
        x = left + int(index * (right - left) / max(len(dates) - 1, 1))
        y = bottom - int((value - low) * (bottom - top) / (high - low))
        return x, y

    for step in range(5):
        value = low + step * (high - low) / 4
        y = point(0, value)[1]
        draw.line((left, y, right, y), fill="#E5E7EB", width=1)
        draw.text((25, y - 14), f"{value:.0f}", fill="#6B7280", font=small)
    draw.line((left, top, left, bottom), fill="#9CA3AF", width=2)
    draw.line((left, bottom, right, bottom), fill="#9CA3AF", width=2)
    draw.line(
        [point(i, value) for i, value in enumerate(byd)], fill=f"#{BLUE}", width=5
    )
    draw.line(
        [point(i, value) for i, value in enumerate(csi300)],
        fill=f"#{GOLD}",
        width=5,
    )
    draw.line((1030, 70, 1080, 70), fill=f"#{BLUE}", width=5)
    draw.text((1090, 55), "BYD", fill="#374151", font=small)
    draw.line((1200, 70, 1250, 70), fill=f"#{GOLD}", width=5)
    draw.text((1260, 55), "CSI 300", fill="#374151", font=small)
    for index in (0, len(dates) // 2, len(dates) - 1):
        x, _ = point(index, low)
        label = dates[index].strftime("%Y-%m")
        box = draw.textbbox((0, 0), label, font=small)
        draw.text((x - (box[2] - box[0]) / 2, 535), label, fill="#6B7280", font=small)
    image.save(market_path, quality=95)

    matrix = Image.new("RGB", (1400, 820), "white")
    draw = ImageDraw.Draw(matrix)
    cn_font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 25)
    cn_small = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 21)
    cn_bold = ImageFont.truetype("C:/Windows/Fonts/msyhbd.ttc", 34)
    draw.text((90, 28), "风险优先级矩阵：影响 × 紧迫度", fill="#0B2545", font=cn_bold)
    left, top, right, bottom = 165, 110, 1290, 700
    cell_w = (right - left) / 5
    cell_h = (bottom - top) / 5
    fills = [
        ["#EAF6EF", "#EAF6EF", "#FFF5D6", "#FFF5D6", "#FDE8E8"],
        ["#EAF6EF", "#FFF5D6", "#FFF5D6", "#FDE8E8", "#FDE8E8"],
        ["#EAF6EF", "#FFF5D6", "#FBE7C6", "#FDE8E8", "#FAD0D0"],
        ["#FFF5D6", "#FFF5D6", "#FDE8E8", "#FAD0D0", "#F4B9B9"],
        ["#FFF5D6", "#FDE8E8", "#FAD0D0", "#F4B9B9", "#EFA2A2"],
    ]
    for urgency in range(1, 6):
        for impact in range(1, 6):
            x0 = left + (urgency - 1) * cell_w
            y0 = bottom - impact * cell_h
            draw.rectangle(
                (x0, y0, x0 + cell_w, y0 + cell_h),
                fill=fills[impact - 1][urgency - 1],
                outline="#FFFFFF",
                width=4,
            )
    risk_points = {
        "BYD-R1": (5, 5),
        "BYD-R2": (5, 5),
        "BYD-R3": (4, 4),
        "BYD-R4": (3, 3),
        "BYD-R5": (5, 4),
        "BYD-R6": (4, 3),
    }
    offsets = {"BYD-R1": (-55, -23), "BYD-R2": (42, 32)}
    for risk_id, (impact, urgency) in risk_points.items():
        x = left + (urgency - 0.5) * cell_w
        y = bottom - (impact - 0.5) * cell_h
        dx, dy = offsets.get(risk_id, (0, 0))
        label = risk_id.replace("BYD-", "")
        box = draw.textbbox((0, 0), label, font=cn_font)
        pad_x, pad_y = 18, 10
        draw.rounded_rectangle(
            (
                x + dx - (box[2] - box[0]) / 2 - pad_x,
                y + dy - (box[3] - box[1]) / 2 - pad_y,
                x + dx + (box[2] - box[0]) / 2 + pad_x,
                y + dy + (box[3] - box[1]) / 2 + pad_y,
            ),
            radius=18,
            fill="#0B2545",
        )
        draw.text(
            (x + dx - (box[2] - box[0]) / 2, y + dy - (box[3] - box[1]) / 2),
            label,
            fill="white",
            font=cn_font,
        )
    for value in range(1, 6):
        draw.text(
            (left + (value - 0.5) * cell_w - 8, bottom + 18),
            str(value),
            fill="#374151",
            font=cn_small,
        )
        draw.text(
            (left - 42, bottom - (value - 0.5) * cell_h - 13),
            str(value),
            fill="#374151",
            font=cn_small,
        )
    draw.text((610, 760), "紧迫度 →", fill="#374151", font=cn_font)
    draw.text((28, 360), "影\n响\n↑", fill="#374151", font=cn_font, spacing=5)
    draw.text((1030, 735), "1=低，5=高", fill="#6B7280", font=cn_small)
    matrix.save(risk_matrix_path, quality=95)
    return finance_path, market_path, risk_matrix_path


def _add_picture(
    doc: Document, path: Path, caption: str, source: str, *, width: float = 6.25
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    _set_run_font(run, size=9, bold=True, color=DARK_GRAY)
    _add_citation(doc, source)


def _risk_color(severity: str) -> str:
    return {"critical": RED, "high": RED, "medium": AMBER, "low": GREEN}[severity]


def _action_execution_fields(index: int) -> tuple[str, str]:
    timelines = [
        "30天内上线；季度复盘",
        "30天内完成；月度更新",
        "4周形成首版；月度滚动",
        "第2季度完成；季度压力测试",
        "90天建树；事件触发更新",
        "投决前启用；月度监测",
        "半年内纳入内控；季度抽样",
        "6个月完成；半年演练",
    ]
    kpis = [
        "经营现金流/净利润；营运资本占用天数",
        "实际担保余额覆盖率=100%；逾期代偿=0",
        "销量下滑原因覆盖率=100%；预测误差≤10%",
        "核心价格/成本情景覆盖率=100%；阈值触发24小时上报",
        "重点司法辖区覆盖率=100%；名单变化48小时完成影响评估",
        "单站利用率；投资回收期；低效项目暂停率",
        "价格对标覆盖率=100%；重大偏差闭环率=100%",
        "关键客户/供应商预案覆盖率≥90%；半年演练1次",
    ]
    return timelines[index], kpis[index]


def build_docx(input_json: Path, output_docx: Path, chart_dir: Path) -> Path:
    artifact = json.loads(input_json.read_text(encoding="utf-8"))
    report = artifact["report"]
    context = artifact["source_context"]
    profile = context["risk_profile"]
    analysis = context["analysis_bundle"]
    debate = context["debate_result"]
    finance_chart, market_chart, risk_matrix_chart = _build_charts(artifact, chart_dir)

    doc = Document()
    _configure_sections(doc)
    _configure_styles(doc)
    bullet_id = _add_numbering(doc, bullet=True)
    doc.core_properties.title = report["title"]
    doc.core_properties.subject = "上市公司公开信息风险与战略咨询"
    doc.core_properties.author = "Financial Research & Advisory Agent"
    doc.core_properties.keywords = "BYD, risk advisory, public information"
    _add_cover(doc, report, artifact)

    doc.add_heading("Partner Executive Summary", level=1)
    _add_callout(
        doc,
        "Bottom line",
        "比亚迪当前并非缺少增长投入，而是增长质量、现金转化与执行节奏出现分化。建议管理层把未来90天聚焦于三件事：解释利润到现金的缺口、定位销量下滑根因、核清集团担保实际暴露；在事实底座补齐前，不宜以全面扩张或全面收缩替代精细化决策。",
    )
    doc.add_heading("需要管理层立即做出的三项决定", level=2)
    decision_num_id = _add_numbering(doc, bullet=False)
    for item in [
        "批准由CFO牵头的现金转化与担保暴露双口径核查，30天内形成管理层底稿。",
        "成立销量与毛利联合工作组，按地区、品牌和渠道建立月度归因与滚动预测。",
        "把地缘合规、闪充投资和关联交易纳入季度风险仪表盘，并明确升级阈值。",
    ]:
        _add_list_item(doc, item, decision_num_id)
    doc.add_heading("关键事实与Partner解读", level=2)
    partner_table = doc.add_table(rows=1, cols=3)
    partner_table.style = "Table Grid"
    _set_table_geometry(partner_table, [2160, 2520, 4680])
    for cell, text in zip(
        partner_table.rows[0].cells,
        ["关键信号", "公开证据", "Partner解读"],
        strict=True,
    ):
        _set_cell_shading(cell, NAVY)
        _set_cell_text(
            cell,
            text,
            bold=True,
            color=WHITE,
            size=9,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    partner_rows = [
        (
            "盈利与现金背离",
            "净利润-18.97%；经营现金流-55.69%",
            "最优先不是再解释收入增长，而是建立利润—现金桥接并明确营运资本责任。",
        ),
        (
            "销量承压",
            "2026年1—4月销量同比-26.02%",
            "必须区分需求、价格、渠道和产品周期；否则资源配置容易误判。",
        ),
        (
            "外部暴露上升",
            "高额担保授权；海外名单与合规不确定性",
            "授权额不等于实际负债、名单不等于制裁，但两者都需要可验证的暴露台账和升级预案。",
        ),
    ]
    for row_values in partner_rows:
        row = partner_table.add_row().cells
        for index, value in enumerate(row_values):
            _set_cell_text(
                row[index],
                value,
                size=8.6,
                bold=index == 0,
                color=NAVY if index == 0 else BLACK,
                align=WD_ALIGN_PARAGRAPH.CENTER
                if index == 1
                else WD_ALIGN_PARAGRAPH.LEFT,
            )
    _set_table_geometry(partner_table, [2160, 2520, 4680])
    _add_callout(
        doc,
        "Partner recommendation",
        "建议采用“先补事实底座、再设触发阈值、最后动态配置资本”的治理顺序。报告中的风险等级用于管理排序，不构成证券评级或信用结论。",
        color=GOLD,
    )
    _add_citation(doc, "证据：BYD-E1、BYD-E2、BYD-E3、BYD-E6；管理建议为本报告判断。")

    _add_section_heading(doc, "1. 执行摘要")
    _add_callout(doc, "总体判断", report["executive_summary"], color=NAVY)
    doc.add_heading("1.1 五项结论先行", level=2)
    for index, finding in enumerate(report["headline_findings"], start=1):
        heading = doc.add_heading(f"{index}. {finding['title']}", level=3)
        _set_paragraph_border(
            heading, _risk_color("high") if index <= 3 else GOLD, size=10, space=5
        )
        _add_body(doc, finding["judgment"])
        _add_body(doc, finding["implication"], bold_lead="管理含义：")
        _add_citation(doc, "证据：" + "、".join(finding["evidence_ids"]))

    doc.add_page_break()
    doc.add_heading("1.2 风险优先级快照", level=2)
    _add_body(
        doc,
        "二维矩阵综合考虑潜在影响与需要管理层介入的时间紧迫度。右上象限应进入月度经营会或专项委员会；其余风险按触发条件升级。",
    )
    _add_picture(
        doc,
        risk_matrix_chart,
        "图1：比亚迪风险优先级二维矩阵",
        "来源：公开证据与本报告管理判断；坐标用于优先级排序，不代表发生概率。",
        width=5.65,
    )
    snapshot = doc.add_table(rows=1, cols=4)
    snapshot.style = "Table Grid"
    _set_table_geometry(snapshot, [1080, 2700, 1260, 4320])
    headers = ["编号", "风险主题", "等级", "需要回答的核心问题"]
    for cell, text in zip(snapshot.rows[0].cells, headers, strict=True):
        _set_cell_shading(cell, NAVY)
        _set_cell_text(
            cell, text, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER
        )
    _set_repeat_table_header(snapshot.rows[0])
    for risk in profile["assessments"]:
        row = snapshot.add_row().cells
        values = [
            risk["risk_id"],
            risk["category"],
            risk["severity"].upper(),
            risk["implication"],
        ]
        for index, (cell, value) in enumerate(zip(row, values, strict=True)):
            _set_cell_text(
                cell,
                value,
                bold=index == 2,
                color=_risk_color(risk["severity"]) if index == 2 else BLACK,
                size=8.2,
                align=WD_ALIGN_PARAGRAPH.CENTER
                if index in {0, 2}
                else WD_ALIGN_PARAGRAPH.LEFT,
            )
    _set_table_geometry(snapshot, [1080, 2700, 1260, 4320])
    _add_citation(
        doc, "来源：风险诊断结构化结果；等级表示监测优先级，不代表信用评级或投资评级。"
    )

    _add_section_heading(doc, "2. 分析范围、方法与证据边界")
    _add_body(doc, report["mandate_and_scope"])
    doc.add_heading("2.1 分析方法", level=2)
    for item in [
        "证据登记：仅使用六份巨潮正式披露及BaoStock行情统计，并保存发布日期、检索日期和证据编号。",
        "风险诊断：从财务、经营、集团信用、治理、执行和地缘合规六个维度识别风险。",
        "Debate Gate：重大咨询任务、推断性结论和资料缺口触发正反方两轮辩论。",
        "咨询写作：DeepSeek依据结构化输入组织报告；证据ID由代码复核，未知ID会阻止报告生成。",
        "人工复核边界：报告不替代管理层内部数据、审计程序、法律意见或投资决策。",
    ]:
        _add_list_item(doc, item, bullet_id)
    doc.add_heading("2.2 关键口径保护", level=2)
    _add_callout(
        doc,
        "担保口径",
        "1,855.15亿元为相关担保授权上限，并非实际担保余额、实际负债或已发生损失。",
        color=RED,
    )
    _add_callout(
        doc,
        "地缘口径",
        "公司披露的美国国防部相关名单并非制裁名单。报告评估的是后续升级和合作方行为变化风险，而非声称已发生全面制裁。",
        color=RED,
    )

    _add_section_heading(doc, "3. 财务、经营与资本市场语境")
    doc.add_heading("3.1 财务与经营信号", level=2)
    _add_body(doc, report["financial_and_operating_context"])
    metrics = doc.add_table(rows=1, cols=3)
    metrics.style = "Table Grid"
    _set_table_geometry(metrics, [3600, 2520, 3240])
    for cell, text in zip(
        metrics.rows[0].cells, ["指标", "2025年变化", "风险解读"], strict=True
    ):
        _set_cell_shading(cell, NAVY)
        _set_cell_text(
            cell, text, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER
        )
    metric_rows = [
        ("营业收入", "+3.46%", "收入仍增长，但不足以单独说明盈利质量"),
        ("归母净利润", "-18.97%", "增长与盈利出现背离"),
        ("经营现金流净额", "-55.69%", "现金转化需要优先拆解"),
        ("加权平均ROE", "26.05% → 15.31%", "资本回报效率下降"),
        ("归母净资产", "+32.94%", "构成缓释能力，但也提高资本效率要求"),
    ]
    for metric, change, meaning in metric_rows:
        row = metrics.add_row().cells
        for index, value in enumerate([metric, change, meaning]):
            _set_cell_text(
                row[index],
                value,
                bold=index == 1,
                size=9,
                align=WD_ALIGN_PARAGRAPH.CENTER
                if index == 1
                else WD_ALIGN_PARAGRAPH.LEFT,
            )
    _set_table_geometry(metrics, [3600, 2520, 3240])
    _add_citation(doc, "来源：BYD-E1。百分比为公司2025年年度报告摘要披露的同比变化。")
    _add_picture(
        doc,
        finance_chart,
        "图2：比亚迪2025年部分财务指标同比变化",
        "来源：BYD-E1；图表由代码根据公开披露绘制。",
    )

    doc.add_heading("3.2 资本市场语境", level=2)
    _add_body(doc, report["market_context"])
    market = context["market"]
    market_table = doc.add_table(rows=1, cols=3)
    market_table.style = "Table Grid"
    _set_table_geometry(market_table, [3960, 2700, 2700])
    for cell, text in zip(
        market_table.rows[0].cells, ["统计指标", "比亚迪", "沪深300"], strict=True
    ):
        _set_cell_shading(cell, NAVY)
        _set_cell_text(
            cell, text, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER
        )
    for label, key in [
        ("2025-01-02以来回报", "period_return_pct"),
        ("2026年初以来回报", "ytd_return_pct"),
        ("年化波动率", "annualized_volatility_pct"),
        ("区间最大回撤", "maximum_drawdown_pct"),
    ]:
        row = market_table.add_row().cells
        _set_cell_text(row[0], label, size=9)
        _set_cell_text(
            row[1],
            f"{market['byd'][key]:.2f}%",
            bold=True,
            size=9,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text(
            row[2],
            f"{market['csi300'][key]:.2f}%",
            size=9,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    _set_table_geometry(market_table, [3960, 2700, 2700])
    _add_citation(
        doc,
        "来源：MARKET-E1，BaoStock前复权日线；行情仅用于风险语境，不用于估值或评级。",
    )
    _add_picture(
        doc,
        market_chart,
        "图3：比亚迪与沪深300相对表现（起点=100）",
        "来源：MARKET-E1，BaoStock前复权日线；图表由代码绘制。",
    )

    _add_section_heading(doc, "4. 六项核心风险诊断")
    risk_by_id = {item["risk_id"]: item for item in profile["assessments"]}
    for index, chapter in enumerate(report["risk_chapters"], start=1):
        risk = risk_by_id[chapter["risk_id"]]
        doc.add_heading(f"4.{index} {chapter['title']}", level=2)
        _add_callout(
            doc,
            f"{chapter['severity'].upper()} PRIORITY",
            risk["observation"],
            color=_risk_color(chapter["severity"]),
        )
        doc.add_heading("诊断", level=3)
        _add_body(doc, chapter["diagnosis"])
        doc.add_heading("为何重要", level=3)
        _add_body(doc, chapter["why_it_matters"])
        doc.add_heading("建议监测的领先指标", level=3)
        for item in chapter["leading_indicators"]:
            _add_list_item(doc, item, bullet_id)
        doc.add_heading("管理层需要回答的问题", level=3)
        question_num_id = _add_numbering(doc, bullet=False)
        for item in chapter["management_questions"]:
            _add_list_item(doc, item, question_num_id)
        doc.add_heading("建议行动", level=3)
        for item in chapter["recommended_actions"]:
            _add_list_item(doc, item, bullet_id)
        _add_body(doc, chapter["residual_uncertainty"], bold_lead="剩余不确定性：")
        _add_citation(doc, "证据：" + "、".join(chapter["evidence_ids"]))

    _add_section_heading(doc, "5. 三种条件情景")
    _add_callout(
        doc,
        "使用说明",
        "情景是条件推演，不是概率预测。其目的在于提前约定可观察触发信号和管理动作。",
        color=GOLD,
    )
    scenario_colors = [GREEN, GOLD, RED]
    for index, (scenario, color) in enumerate(
        zip(report["scenarios"], scenario_colors, strict=True), start=1
    ):
        heading = doc.add_heading(f"5.{index} {scenario['name']}", level=2)
        _set_paragraph_border(heading, color, size=14, space=6)
        _add_body(doc, scenario["narrative"])
        doc.add_heading("可观察触发信号", level=3)
        for item in scenario["observable_triggers"]:
            _add_list_item(doc, item, bullet_id)
        doc.add_heading("业务影响", level=3)
        for item in scenario["business_implications"]:
            _add_list_item(doc, item, bullet_id)
        doc.add_heading("建议管理响应", level=3)
        response_num_id = _add_numbering(doc, bullet=False)
        for item in scenario["management_response"]:
            _add_list_item(doc, item, response_num_id)

    _add_section_heading(doc, "6. 研究辩论：共识、争议与资料缺口")
    _add_body(doc, report["debate_synthesis"])
    doc.add_heading("6.1 已形成共识", level=2)
    for item in debate["consensus_findings"]:
        _add_list_item(doc, item, bullet_id)
    doc.add_heading("6.2 仍有争议", level=2)
    for item in debate["disputed_findings"]:
        _add_list_item(doc, item, bullet_id)
    doc.add_heading("6.3 不能由模型替代的数据", level=2)
    for item in debate["unresolved_issues"]:
        _add_list_item(doc, item, bullet_id)
    _add_callout(doc, "Moderator结论", debate["moderator_conclusion"], color=NAVY)

    _add_section_heading(doc, "7. 0-12个月行动路线图")
    _add_body(
        doc,
        "行动建议按信息获取、风险量化和管理响应三个层次展开。前3个月优先消除关键数据盲区，随后开展压力测试与治理机制建设。",
    )
    doc.add_heading("7.1 执行责任表", level=2)
    roadmap = doc.add_table(rows=1, cols=5)
    roadmap.style = "Table Grid"
    roadmap_widths = [540, 3150, 1440, 1350, 2880]
    _set_table_geometry(roadmap, roadmap_widths)
    for cell, text in zip(
        roadmap.rows[0].cells,
        ["编号", "行动", "Owner", "Timeline", "KPI / 完成标准"],
        strict=True,
    ):
        _set_cell_shading(cell, NAVY)
        _set_cell_text(
            cell,
            text,
            bold=True,
            color=WHITE,
            size=8.4,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for index, action in enumerate(report["priority_actions"]):
        timeline, kpi = _action_execution_fields(index)
        row = roadmap.add_row().cells
        values = [
            f"A{index + 1}",
            action["action"],
            action["proposed_owner"],
            timeline,
            kpi,
        ]
        for column, value in enumerate(values):
            _set_cell_text(
                row[column],
                value,
                size=8.0,
                bold=column == 0,
                color=NAVY if column == 0 else BLACK,
                align=WD_ALIGN_PARAGRAPH.CENTER
                if column in {0, 2, 3}
                else WD_ALIGN_PARAGRAPH.LEFT,
            )
        if action["horizon"] == "0-3个月":
            _set_cell_shading(row[0], "FDE8E8")
        elif action["horizon"] == "3-6个月":
            _set_cell_shading(row[0], "FFF5D6")
        else:
            _set_cell_shading(row[0], "EAF6EF")
    _set_table_geometry(roadmap, roadmap_widths)
    _add_citation(
        doc,
        "注：Owner为建议牵头部门；Timeline与KPI为本报告提出的执行标准，需由管理层确认并纳入正式绩效机制。",
    )

    _add_section_heading(doc, "8. 风险监测仪表盘")
    dashboard = doc.add_table(rows=1, cols=4)
    dashboard.style = "Table Grid"
    _set_table_geometry(dashboard, [2520, 3600, 1440, 1800])
    for cell, text in zip(
        dashboard.rows[0].cells,
        ["风险主题", "核心监测指标", "建议频率", "升级条件"],
        strict=True,
    ):
        _set_cell_shading(cell, NAVY)
        _set_cell_text(
            cell,
            text,
            bold=True,
            color=WHITE,
            size=8.8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    frequencies = ["季度", "月度", "月度", "季度", "持续", "月度"]
    escalation = [
        "现金转化继续恶化",
        "销量与利润同步承压",
        "余额或代偿风险上升",
        "定价或集中度异常",
        "限制范围或司法辖区扩大",
        "进度偏离且利用率不足",
    ]
    for risk, frequency, trigger in zip(
        profile["assessments"], frequencies, escalation, strict=True
    ):
        row = dashboard.add_row().cells
        values = [
            risk["category"],
            "；".join(risk["monitoring_indicators"]),
            frequency,
            trigger,
        ]
        for index, value in enumerate(values):
            _set_cell_text(
                row[index],
                value,
                size=8.4,
                align=WD_ALIGN_PARAGRAPH.CENTER
                if index in {2, 3}
                else WD_ALIGN_PARAGRAPH.LEFT,
            )
    _set_table_geometry(dashboard, [2520, 3600, 1440, 1800])
    _add_citation(doc, "注：频率与升级条件为本报告提出的管理建议，不是公司已披露制度。")

    _add_section_heading(doc, "9. 结论")
    _add_callout(doc, "结论", report["conclusion"], color=NAVY)
    _add_body(
        doc,
        "本报告的价值不在于给出单一风险分数，而在于明确管理层现在需要知道什么、哪些口径不能误读、哪些数据必须补充，以及在不同条件下应采取什么行动。",
    )

    _add_section_heading(doc, "附录A：公开证据索引")
    evidence_table = doc.add_table(rows=1, cols=4)
    evidence_table.style = "Table Grid"
    _set_table_geometry(evidence_table, [1260, 3600, 1620, 2880])
    for cell, text in zip(
        evidence_table.rows[0].cells, ["编号", "文件", "发布日期", "用途"], strict=True
    ):
        _set_cell_shading(cell, NAVY)
        _set_cell_text(
            cell,
            text,
            bold=True,
            color=WHITE,
            size=8.8,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    risk_evidence = {
        evidence_id: risk["category"]
        for risk in profile["assessments"]
        for evidence_id in risk["evidence_ids"]
    }
    for evidence in analysis["evidence"]:
        row = evidence_table.add_row().cells
        published = (evidence.get("published_at") or "未知")[:10]
        values = [
            evidence["evidence_id"],
            evidence["title"],
            published,
            risk_evidence.get(evidence["evidence_id"], "背景与缓释因素"),
        ]
        for index, value in enumerate(values):
            _set_cell_text(
                row[index],
                value,
                size=8.3,
                align=WD_ALIGN_PARAGRAPH.CENTER
                if index in {0, 2}
                else WD_ALIGN_PARAGRAPH.LEFT,
            )
    _set_table_geometry(evidence_table, [1260, 3600, 1620, 2880])
    for evidence in analysis["evidence"]:
        _add_citation(
            doc, f"{evidence['evidence_id']}｜{evidence['url']}｜{evidence['summary']}"
        )
    _add_citation(
        doc,
        "MARKET-E1｜BaoStock前复权日线，2025-01-02至2026-08-07，检索时间见JSON审计产物。",
    )

    _add_section_heading(doc, "附录B：限制与免责声明")
    for item in profile["scope_limitations"]:
        _add_list_item(doc, item, bullet_id)
    for item in [
        "DeepSeek仅负责将结构化证据组织为咨询式文字；财务数字、行情统计、证据ID和辩论轮数由代码生成或校验。",
        "报告没有接入公司非公开经营数据、合同明细、海外法律意见、客户访谈或供应链尽调。",
        "风险等级代表管理关注优先级，不构成证券评级、信用评级或违约概率。",
        "随着新公告、财务报告和监管变化出现，应重新运行证据收集、Debate Gate与报告流程。",
    ]:
        _add_list_item(doc, item, bullet_id)

    for table in doc.tables:
        if table.rows:
            _set_repeat_table_header(table.rows[0])
        for row in table.rows:
            _set_row_cant_split(row)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx.resolve()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--chart-dir", type=Path, required=True)
    args = parser.parse_args()
    path = build_docx(args.input, args.output, args.chart_dir)
    print(path)


if __name__ == "__main__":
    main()
